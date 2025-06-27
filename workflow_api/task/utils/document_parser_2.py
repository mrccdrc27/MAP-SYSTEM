import json
import re
import os
from datetime import datetime

import pdfplumber
from dateutil import parser as date_parser  # pip install python-dateutil
from docx import Document  # pip install python-docx

# -------------------- TEXT EXTRACTION --------------------

def extract_text_from_pdf(file_path):
    """
    Extracts and concatenates text from all pages of a PDF file.
    """
    with pdfplumber.open(file_path) as pdf:
        texts = []
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                texts.append(txt)
        return "\n".join(texts)


def extract_text_from_docx(file_path):
    """
    Extracts and concatenates text from all paragraphs of a DOCX file.
    """
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)

# -------------------- TABLE EXTRACTION --------------------

def format_cost(cost_str):
    """
    Normalize cost strings to a two-decimal format.
    """
    try:
        return "{:.2f}".format(float(cost_str.replace(',', '').replace('$', '')))
    except Exception:
        return cost_str.strip()


def extract_items_from_docx_table(doc):
    """
    Extracts item rows from any DOCX tables matching BMS headers.
    """
    items = []
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "cost element" in headers and "description" in headers:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 4:
                    item = {
                        "cost_element": cells[0],
                        "description": cells[1],
                        "estimated_cost": format_cost(cells[2]),
                        "account": int(cells[3]) if cells[3].isdigit() else cells[3],
                    }
                    if len(cells) > 4 and cells[4]:
                        item["notes"] = cells[4]
                    items.append(item)
    return items


def extract_items_from_pdf_table(file_path):
    """
    Extracts item rows from any PDF tables matching BMS headers.
    """
    items = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                headers = [cell.strip().lower() if cell else '' for cell in table[0]]
                if "cost element" in headers and "description" in headers:
                    for row in table[1:]:
                        cells = [c.replace("\n", " ").strip() if c else '' for c in row]
                        if len(cells) >= 4:
                            item = {
                                "cost_element": cells[0],
                                "description": cells[1],
                                "estimated_cost": format_cost(cells[2]),
                                "account": int(cells[3]) if cells[3].isdigit() else cells[3],
                            }
                            if len(cells) > 4 and cells[4]:
                                item["notes"] = cells[4]
                            items.append(item)
    return items

# -------------------- DATE NORMALIZATION --------------------

def normalize_date(value):
    """
    Parse any date-like string into ISO 8601 date (YYYY-MM-DD).
    """
    try:
        return date_parser.parse(value).date().isoformat()
    except Exception:
        return None

# -------------------- STRUCTURED FIELD PARSER (BMS) --------------------

def parse_structured_fields(text):
    """
    Extracts BMS fields from text using regex patterns, defaulting missing fields to None.
    """
    keys = [
        "title", "project_summary", "project_description", "performance_notes",
         "department_input", "fiscal_year", "submitted_by_name", "status",
         "performance_start_date", "performance_end_date", "ticket_id"
    ]
    result = {k: None for k in keys}

    patterns = {
        "title": r"(?:Title|title):\s*(.+)",
        "project_summary": r"(?:Project Summary|project_summary):\s*(.+)",
        "project_description": r"(?:Project Description|project_description):\s*(.+)",
        "performance_notes": r"(?:Performance Notes|performance_notes):\s*(.+)",
        "department_input": r"(?:Department Input|department_input|Department):\s*(\d+)",
        "fiscal_year": r"(?:Fiscal Year|fiscal_year):\s*(\d+)",
        "submitted_by_name": r"(?:Submitted By|submitted_by_name):\s*(.+)",
        "status": r"(?:Status|status):\s*(\w+)",
        "performance_start_date": r"(?:Performance Start Date|performance_start_date):\s*(.+)",
        "performance_end_date": r"(?:Performance End Date|performance_end_date):\s*(.+)",
        "ticket_id": r"(?:Ticket ID|ticket_id):\s*([A-Z]+-[A-Z]+-\d{4}-\d+)"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if key in ["performance_start_date", "performance_end_date"]:
                result[key] = normalize_date(value)
            elif key in ["department_input", "fiscal_year"]:
                result[key] = int(value)
            else:
                result[key] = value
    return result

# -------------------- AMS FIELD PARSER --------------------

def parse_ams_fields(text):
    """
    Extracts AMS-specific fields from text using regex patterns, defaulting missing fields to None.
    """
    keys = [
        "ticket_id", "asset_id", "asset_name", "requestor_id", "requestor",
        "requestor_location", "checkout_date", "return_date", "checkin_date",
        "checkout_ref_id", "condition", "is_resolved"
    ]
    result = {k: None for k in keys}

    patterns = {
        # "ticket_id": r"(?:Ticket ID|ticket_id):\s*([\w-]+)",
        "asset_id": r"(?:Asset ID|asset_id):\s*(\d+)",
        "asset_name": r"(?:Asset Name|asset_name):\s*(.+)",
        "requestor_id": r"(?:Requestor ID|requestor_id):\s*(\d+)",
        "requestor": r"(?:Requestor|requestor):\s*(.+)",
        "requestor_location": r"(?:Requestor Location|requestor_location):\s*(.+)",
        "checkout_date": r"(?:Checkout Date|checkout_date):\s*(.+)",
        "return_date": r"(?:Return Date|return_date):\s*(.+)",
        "checkin_date": r"(?:Checkin Date|checkin_date):\s*(.+)",
        "checkout_ref_id": r"(?:Checkout Ref ID|checkout_ref_id):\s*(.+)",
        "condition": r"(?:Condition|condition):\s*(.+)",
        "is_resolved": r"(?:Is Resolved|is_resolved):\s*(true|false)"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if key in ["asset_id", "requestor_id"]:
                result[key] = int(value)
            elif key in ["checkout_date", "return_date", "checkin_date"]:
                result[key] = normalize_date(value)
            elif key == "is_resolved":
                result[key] = value.lower() == "true"
            else:
                result[key] = value
    return result

# -------------------- DOCUMENT TYPE DETECTION --------------------

def detect_document_type(text):
    """
    Detects if the document is BMS or AMS based on a header line.
    """
    match = re.search(r"Document Type:\s*(\w+)", text, re.IGNORECASE)
    return match.group(1).strip().upper() if match else None

# -------------------- MAIN PROCESSOR --------------------

def process_document(file_path):
    """
    Main entry point: reads file, detects type, dispatches to appropriate parser,
    and bundles items if BMS. Ensures missing fields are null.
    Returns a JSON string of the extracted data.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        doc = None
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
        doc = Document(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    doc_type = detect_document_type(text)
    if doc_type == "BMS":
        data = parse_structured_fields(text)
        items = extract_items_from_pdf_table(file_path) if ext == ".pdf" else extract_items_from_docx_table(doc)
        data["items"] = items if items else []
    elif doc_type == "AMS":
        data = parse_ams_fields(text)
    else:
        raise ValueError("Unsupported or missing Document Type header.")

    return json.dumps(data, indent=2)

# -------------------- SCRIPT TEST --------------------

if __name__ == "__main__":
    test_files = [
        r"C:\work\god\Ticket-Tracking-System\workflow_api\task\utils\BMSDOCU.docx",
    ]
    for file_path in test_files:
        print(f"\nProcessing: {file_path}")
        try:
            output = process_document(file_path)
            print("Resulting JSON:\n", output)
        except Exception as e:
            print("Error:", e)
