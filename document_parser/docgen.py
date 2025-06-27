from docx import Document
from faker import Faker
import random
import os
from datetime import datetime, timedelta

# Setup
fake = Faker()
output_dir = "generated_docs"
os.makedirs(output_dir, exist_ok=True)

# Configuration
NUM_DOCUMENTS = 5
MIN_TABLE_ROWS = 2
MAX_TABLE_ROWS = 6

def generate_bms_document(doc_index):
    doc = Document()

    # Header
    doc.add_heading('Document Type: BMS', level=1)
    doc.add_paragraph('')

    # Generate Dates
    start_date = fake.date_between(
        start_date=datetime.strptime("2026-01-01", "%Y-%m-%d").date(),
        end_date=datetime.strptime("2026-01-31", "%Y-%m-%d").date()
    )
    end_date = start_date + timedelta(days=random.randint(30, 90))

    # Metadata fields
    fields = {
        "Performance Start Date": start_date.strftime("%Y-%m-%d"),
        "Performance End Date": end_date.strftime("%Y-%m-%d"),
        "Ticket ID": f"TTF-{fake.lexify('???').upper()}-{start_date.year}-00{doc_index+1}",
        "Status": "SUBMITTED",
        "Submitted By": "External System Interface",
        "Title": fake.catch_phrase(),
        "Project Summary": fake.sentence(nb_words=6),
        "Project Description": fake.paragraph(nb_sentences=2),
        "Performance Notes": fake.sentence(),
        "Department Input": str(random.randint(1, 5)),
        "Fiscal Year": 2
    }

    for key, value in fields.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_paragraph('')

    # Cost Table Header
    headers = ["Cost Element", "Description", "Estimated Cost", "Account"]
    num_rows = random.randint(MIN_TABLE_ROWS, MAX_TABLE_ROWS)

    cost_elements = ["Software License", "Training", "Support", "Maintenance", "Consulting"]
    accounts = ["1", "2", "3"]

    # Add Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for _ in range(num_rows):
        cost_element = random.choice(cost_elements)
        description = fake.sentence(nb_words=5)
        cost = str(random.randint(2000, 50000))
        account = random.choice(accounts)

        row = table.add_row().cells
        row[0].text = cost_element
        row[1].text = description
        row[2].text = cost
        row[3].text = account

    # Save
    filename = f"BMS_Ticket_{doc_index+1:03d}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"✅ Generated: {filepath}")

# Main loop
for i in range(NUM_DOCUMENTS):
    generate_bms_document(i)
