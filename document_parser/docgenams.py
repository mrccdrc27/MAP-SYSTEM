from docx import Document
import os

# Input data from docgenams.py
data = [
    {
        "asset_id": 22,
        "asset_name": "iPad Pro",
        "requestor_id": 1,
        "requestor": "Kelly Barron",
        "requestor_location": "Pasig Office",
        "checkout_date": "2025-06-16",
        "return_date": "2025-08-16",
        "checkin_date": None,
        "checkout_ref_id": None,
        "condition": None,
    },
    {
        "asset_id": 24,
        "asset_name": "LENOVO",
        "requestor_id": 2,
        "requestor": "Tammy Sawyer",
        "requestor_location": "Quezon City Office",
        "checkout_date": "2025-06-13",
        "return_date": "2028-06-16",
        "checkin_date": None,
        "checkout_ref_id": None,
        "condition": None,
    },
    {
        "asset_id": 25,
        "asset_name": "Samsung",
        "requestor_id": 3,
        "requestor": "Denise Jimenez",
        "requestor_location": "Manila Office",
        "checkout_date": "2025-06-13",
        "return_date": "2027-06-16",
        "checkin_date": "2027-06-16",
        "checkout_ref_id": 1,
        "condition": 9,
    },
    {
        "ticket_id": "TK-6018",
        "asset_id": 27,
        "asset_name": "Samsuung 12",
        "requestor_id": 3,
        "requestor": "Denise Jimenez",
        "requestor_location": "Makati Office",
        "checkout_date": "2025-06-10",
        "return_date": "2025-06-16",
        "checkin_date": "2025-06-18",
        "checkout_ref_id": 2,
        "condition": 6,
    },
]

# Output directory
output_dir = "generated_docs"
os.makedirs(output_dir, exist_ok=True)

def generate_ams_document(record, doc_index):
    doc = Document()

    # Header
    doc.add_heading('Document Type: AMS', level=1)
    doc.add_paragraph('')

    # Add fields to the document
    for key, value in record.items():
        doc.add_paragraph(f"{key}: {value if value is not None else 'N/A'}")

    # Save the document
    filename = f"AMS_Ticket_{doc_index+1:03d}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"✅ Generated: {filepath}")

# Generate documents
for i, record in enumerate(data):
    generate_ams_document(record, i)
