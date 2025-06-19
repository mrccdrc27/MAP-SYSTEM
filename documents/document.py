import json
import re
import os
from docx import Document
import pdfplumber


# -------------------- TEXT EXTRACTION --------------------

def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)


# -------------------- TABLE EXTRACTION --------------------

def extract_items_from_docx_table(doc):
    items = []
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "cost element" in headers and "description" in headers:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 5:
                    item = {
                        "cost_element": cells[0],
                        "description": cells[1],
                        "estimated_cost": cells[2],
                        "account": int(cells[3]) if cells[3].isdigit() else cells[3],
                        "notes": cells[4],
                    }
                    items.append(item)
    return items

def extract_items_from_pdf_table(file_path):
    items = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                headers = [cell.strip().lower() if cell else '' for cell in table[0]]
                if "cost element" in headers and "description" in headers:
                    for row in table[1:]:
                        cells = [cell.replace("\n", " ").strip() if cell else '' for cell in row]
                        if len(cells) >= 5:
                            item = {
                                "cost_element": cells[0],
                                "description": cells[1],
                                "estimated_cost": cells[2],
                                "account": int(cells[3]) if cells[3].isdigit() else cells[3],
                                "notes": cells[4],
                            }
                            items.append(item)
    return items


# -------------------- TEXT FIELD PARSER --------------------

def parse_structured_fields(text):
    patterns = {
        "name": r"Name:\s*(.+)",
        "email": r"Email:\s*(\S+@\S+)",
        "date": r"Date:\s*(\d{2}/\d{2}/\d{4})",
        "reference": r"Reference ID:\s*(\w+)",
        "asset": r"Asset:\s*(.+)",
        "requestor": r"Requestor:\s*(.+)",
        "requestor_location": r"Requestor Location:\s*(.+)",
        "checkout_date": r"Checkout Date:\s*(\d{4}-\d{2}-\d{2})",
        "return_date": r"Return Date:\s*(\d{4}-\d{2}-\d{2}|None|null)?",
        "title": r"Title:\s*(.+)",
        "project_summary": r"Project Summary:\s*(.+)",
        "project_description": r"Project Description:\s*(.+)",
        "performance_notes": r"Performance Notes:\s*(.+)",
        "department": r"Department:\s*(\d+)",
        "fiscal_year": r"Fiscal Year:\s*(\d{4})",
        "submitted_by_name": r"Submitted By:\s*(.+)",
        "status": r"Status:\s*(\w+)",
        "performance_start_date": r"Performance Start Date:\s*(\d{4}-\d{2}-\d{2})",
        "performance_end_date": r"Performance End Date:\s*(\d{4}-\d{2}-\d{2})",
        "external_system_id": r"External System ID:\s*(\w+-\d+)",
    }

    result = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result[key] = match.group(1)
    return result


# -------------------- MAIN --------------------

def process_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        items = extract_items_from_pdf_table(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
        doc = Document(file_path)
        items = extract_items_from_docx_table(doc)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    data = parse_structured_fields(text)

    if items:
        data["items"] = items

    return json.dumps(data, indent=2)


# -------------------- TEST --------------------

if __name__ == "__main__":
    file_path = r"C:\work\Ticket-Tracking-System\test\Name.pdf"  # or .pdf
    print("📄 Processing:", file_path)
    try:
        result = process_document(file_path)
        print("✅ Extracted JSON:\n", result)
    except Exception as e:
        print("❌ Error:", e)
